import pandas as pd
import re
from docx import Document
from docx.shared import RGBColor, Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# *********************************************************************** Modules ********************************************************************************************************
def preprocess_llm_answer(text):
    try:
        pattern1 = r'\|\s*(AC-\d+)\s*\|\s*(IS-\d+)\s*\|'
        pattern2 = r'\|\s*Action ID\s*\|\s*Issue ID\s*\|\n\|[-\s]+\|\n\|((?:\s*AC-\d+\s*\|\s*IS-\d+\s*\|\n?)+)'
        
        matches1 = re.findall(pattern1, text)
        matches2 = re.search(pattern2, text)
        
        if matches2:
            matches1 = re.findall(pattern1, matches2.group(1))
        
        if not matches1:
            return text
        
        formatted_lines = [f'{i + 1}. Action ID: {match[0]}, Issue ID: {match[1]}' for i, match in enumerate(matches1)]
        return '\n'.join(formatted_lines)
    
    except Exception as e:
        print(f"Error processing LLM Answer: {str(e)}")
        return text  # Return original text in case of any error

def convert_to_string(value):
    if pd.notnull(value):
        return str(value)
    else:
        return ''

def mapping_df(df, p1, p2, p3):
    try:
        new_df = pd.DataFrame(columns=['Issue ID', 'Issue Name', 'Description', 'Root Cause Description',
                                       'Rating Rationale', 'Risk Category', 'Issue Rating', 'Repeat Issue',
                                       'Status', 'Issue Target Date'])

        issue_id_pattern = re.compile(r'^(?:\d+\.\s+)?-?\s*(?:Issue\s+)?ID:\s+(IS-\d+)')
        issue_name_pattern = re.compile(r'^-?\s*(?:Issue\s+)?Name:\s+(.*)')
        description_pattern = re.compile(r'^-?\s*(?:Issue\s+-\s+)?Description:\s+(.*)')
        root_cause_pattern = re.compile(r'^Root Cause Description:\s+(.*)')
        rating_rationale_pattern = re.compile(r'^(?:Issue\s+)?Rating Rationale:\s+(.*)$')
        risk_category_pattern = re.compile(r'^-?\s*Risk Category:\s*(.*)$')
        issue_rating_pattern = re.compile(r'^-?\s*Issue Rating:\s+(.*)')
        repeat_issue_pattern = re.compile(r'^-?\s*Repeat Issue:\s+(.*)')
        status_pattern = re.compile(r'^-?\s*(?:Issue\s+)?Status:\s+(.*)')
        issue_target_date_pattern = re.compile(r'^-?\s*(?:Issue\s+)?Target Date:\s+(.*)')

        exclude_pattern = re.compile(r'are as follows:|are following result:|Here are the details of the action plans mentioned in the document:|The issues are as follows:')
        current_issue = {}
        Issue_flg = None

        for index, row in df.iterrows():
            if row['Section'] == 'Issue Details':
                Issue_flg = 1
            else:
                pass

            if row[p1] and row[p1].strip() and row[p1] != 'Issue Details':
                Issue_flg = None

            if Issue_flg == 1:
                lines = row['LLM Answer'].split('\n')
                for line in lines:
                    line = line.strip()
                    if exclude_pattern.search(line):
                        continue
                    if issue_id_pattern.match(line):
                        if current_issue:
                            new_df.loc[len(new_df)] = current_issue
                            current_issue = {}
                        current_issue['Issue ID'] = issue_id_pattern.match(line).group(1)
                    elif issue_name_pattern.match(line):
                        current_issue['Issue Name'] = issue_name_pattern.match(line).group(1)
                    elif description_pattern.match(line):
                        current_issue['Description'] = description_pattern.match(line).group(1)
                    elif root_cause_pattern.match(line):
                        current_issue['Root Cause Description'] = root_cause_pattern.match(line).group(1)
                    elif rating_rationale_pattern.match(line):
                        current_issue['Rating Rationale'] = rating_rationale_pattern.match(line).group(1)
                    elif risk_category_pattern.match(line):
                        current_issue['Risk Category'] = risk_category_pattern.match(line).group(1)
                    elif issue_rating_pattern.match(line):
                        current_issue['Issue Rating'] = issue_rating_pattern.match(line).group(1)
                    elif repeat_issue_pattern.match(line):
                        current_issue['Repeat Issue'] = repeat_issue_pattern.match(line).group(1)
                    elif status_pattern.match(line):
                        current_issue['Status'] = status_pattern.match(line).group(1)
                    elif issue_target_date_pattern.match(line):
                        current_issue['Issue Target Date'] = issue_target_date_pattern.match(line).group(1)

        if current_issue:
            new_df.loc[len(new_df)] = current_issue

        return new_df

    except Exception as e:
        print(f"Error in mapping_df: {str(e)}")
        return pd.DataFrame()  # Return empty DataFrame or None as appropriate

def mapp_action_df(df, p1, p2, p3):
    try:
        new_df = pd.DataFrame(columns=[
            'Action Plan ID', 'Action Plan Description', 'Action Plan Owner', 'Action Plan Closure Target Date', 'Action Plan Status'
        ])

        action_plan_id_pattern = re.compile(r'^-?\s*(?:\d+\.\s+Action\s+Plan\s+)?ID:\s+(AC-\d+)$')
        description_pattern = re.compile(r'^-?\s*Description:\s+(.*)')
        owner_pattern = re.compile(r'^-?\s*Action Plan Owner\.Display Name:\s+(.*)')
        due_date_pattern = re.compile(r'^-?\s*(?:Due|Target)\s+[Dd]ate:\s+(.*)$', re.IGNORECASE)
        status_pattern = re.compile(r'^-?\s*Status:\s+(.*)')
        issue_id_pattern = re.compile(r'^-?\s*Issue ID:\s+(IS-\d+)')

        exclude_pattern = re.compile(r'are as follows:|are following result:|Here are the details of the action plans mentioned in the document:')
        current_record = {}
        record_flag = False

        for index, row in df.iterrows():
            if row['Section'] == 'Action Plan Details':
                record_flag = True
            else:
                record_flag = False

            if row[p1] and row[p1].strip() and row[p1] != 'Action Plan Details':
                record_flag = False

            if record_flag:
                lines = row['LLM Answer'].split('\n')
                for line in lines:
                    line = line.strip()
                    if exclude_pattern.search(line):
                        continue
                    if action_plan_id_pattern.match(line):
                        if current_record:
                            new_df.loc[len(new_df)] = current_record
                            current_record = {}
                        current_record['Action Plan ID'] = action_plan_id_pattern.match(line).group(1)
                    elif description_pattern.match(line):
                        current_record['Action Plan Description'] = description_pattern.match(line).group(1)
                    elif owner_pattern.match(line):
                        current_record['Action Plan Owner'] = owner_pattern.match(line).group(1)
                    elif due_date_pattern.match(line):
                        current_record['Action Plan Closure Target Date'] = due_date_pattern.match(line).group(1)
                    elif status_pattern.match(line):
                        current_record['Action Plan Status'] = status_pattern.match(line).group(1)

        if current_record:
            new_df.loc[len(new_df)] = current_record

        return new_df
    
    except Exception as e:
        print(f"Error in mapp_action_df: {str(e)}")
        return pd.DataFrame()  # Return empty DataFrame or None as appropriate

def combine_issues_by_id(df):
    # Define a custom aggregation function that ignores NaN values and joins non-empty strings
    def join_nonempty_strings(series):
        return ' '.join(filter(None, series.dropna().astype(str)))
    
    # Group by 'Issue ID' and concatenate the specified fields
    df_grouped = df.groupby('Issue ID').agg({
        'Issue Name': 'first',  # Take the first 'Issue Name' for each 'Issue ID'
        'Description': join_nonempty_strings,  # Concatenate the 'Description' field
        'Root Cause Description': join_nonempty_strings, 
        'Rating Rationale': join_nonempty_strings,  
        'Risk Category': join_nonempty_strings,  
        'Issue Rating': join_nonempty_strings, 
        'Repeat Issue': join_nonempty_strings,  
        'Status': join_nonempty_strings,  
        'Issue Target Date': join_nonempty_strings 
    }).reset_index()
    
    return df_grouped

def map_dataframes(action_plan_df, issue_df, is_map):
    try:
        # Parse the is_map list to create a mapping dictionary
        mapping_dict = {}
        for mapping in is_map:
            parts = mapping.split(', ')
            action_id = parts[0].split(': ')[1]
            issue_id = parts[1].split(': ')[1]
            mapping_dict[action_id] = issue_id

        # Add a new column to the action_plan_df for the Issue ID
        action_plan_df['Issue ID'] = action_plan_df['Action Plan ID'].map(mapping_dict)

        # Merge the action_plan_df with the issue_df on the Issue ID
        merged_df = action_plan_df.merge(issue_df, on='Issue ID', how='left')

        # Reorder the columns to have Issue ID and related columns first
        issue_columns = [col for col in issue_df.columns if col in merged_df.columns]
        action_plan_columns = [col for col in action_plan_df.columns if col not in issue_columns]
        
        # Ensure 'Issue ID' is the first column
        issue_columns.insert(0, issue_columns.pop(issue_columns.index('Issue ID')))
        merged_df = merged_df[issue_columns + action_plan_columns]

        return merged_df

    except KeyError as e:
        print(f"KeyError occurred: {e}")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        # Handle other unexpected exceptions as needed
        return None



def shade_paragraph(paragraph, shade_color):
    """
    Apply a background shade to a paragraph.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), shade_color)
    paragraph._p.get_or_add_pPr().append(shading_elm)

def add_border(paragraph):
    """
    Add an outside border to a paragraph.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_elm = OxmlElement(f'w:{border_name}')
        border_elm.set(qn('w:val'), 'single')
        border_elm.set(qn('w:sz'), '4')  # Border size (in eighths of a point)
        border_elm.set(qn('w:space'), '1')  # The space between the border and text (in points)
        border_elm.set(qn('w:color'), '000000')  # Border color (black)
        pBdr.append(border_elm)
    pPr.append(pBdr)


def add_dataframe_to_doc(dataframe, doc):
    try:
        # Define the color for the column names
        heading_color = RGBColor(0x1F, 0x4E, 0x79)  # Adjust the RGB values as needed
        # Initialize a variable to store the previous row values
        previous_row_values = {}
        shade_color = 'A6B1B8'  # Approximation of Blue-Grey, Text2, Lighter 40%
        lines_per_page = 10  # Adjust this based on your document's actual lines per page
        line_count = 0
        # Find the index of 'Action Plan ID' column, if it exists
        action_plan_id_index = dataframe.columns.get_loc('Action Plan ID') if 'Action Plan ID' in dataframe.columns else -1
        
        # Process the DataFrame
        for index, row in dataframe.iterrows():
            for col_index, col_name in enumerate(dataframe.columns):
                # Check if 'Issue ID' exists in the row
                if 'Issue ID' in row:
                    # Check if the 'Issue ID' has changed, if so, reset the previous_row_values
                    if index > 0 and row['Issue ID'] != previous_row_values.get('Issue ID'):
                        previous_row_values = {}

                # Check if 'Action Plan ID' exists in the row
                if 'Action Plan ID' in row:
                    # Check if 'Action Plan ID' is about to be printed at the end of the page
                    # import pdb; pdb.set_trace()
                    if col_name == 'Action Plan ID' and line_count >= lines_per_page - 1:
                        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                        line_count = 0
                
                # Check if the current column is before 'Action Plan ID'
                if col_index < action_plan_id_index:
                    # If the value is the same as the previous row, skip it
                    if index > 0 and row[col_name] == previous_row_values.get(col_name):
                        continue
                    
                # Update the previous row value for the current column
                previous_row_values[col_name] = row[col_name]
                
                # Add the value to the document
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{col_name}: ")
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = heading_color  # Apply the color
                run = paragraph.add_run(str(row[col_name]))
                line_count += 1
                
                # Apply a shaded background to specific columns
                if col_name in ['Issue ID', 'Issue Name', 'Action Plan ID', 'Action Plan Description']:
                    shade_paragraph(paragraph, shade_color)  # Light gray background
                
                # Add an outside border to specific columns
                if col_name in ['Issue ID', 'Issue Name', 'Action Plan ID', 'Action Plan Description']:
                    add_border(paragraph)
            
            # Add a line break after each row
            doc.add_paragraph()
            line_count += 1  # Increment line count for the line break
        
        return True
    
    except KeyError as e:
        print(f"KeyError occurred: {e}")
        # Handle KeyError appropriately (e.g., log, raise custom exception, return False)
        return False
    
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        # Handle other exceptions as needed (e.g., log, raise custom exception, return False)
        return False

def create_combined_audit_report(df, p1, p2, p3, header_image_path, merged_dataframe,output_path):
    try:
        doc = Document()
        # Set the default style of the document to Arial
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        # Define the color for the headings
        heading_color = RGBColor(0x1F, 0x4E, 0x79)  # Adjust the RGB values as needed
        
        # Add header image
        section = doc.sections[0]
        section.top_margin = Cm(1.27)  # Adjust the header from top to 1.27 cm
        section.bottom_margin = Cm(0.76)  # Adjust the footer from bottom to 0.76 cm
        header = section.header
        header_paragraph = header.add_paragraph()
        run = header_paragraph.add_run()
        run.add_picture(header_image_path, width=Inches(6))
        header_paragraph.paragraph_format.space_before = Cm(0.5)
        spacing_after = Pt(6)
        
        # Add the main content from the DataFrame
        exclude_pattern = re.compile(r'are as follows:|are following result:|Here are the details of the action plans mentioned in the document:')
        
        # Loop through the DataFrame and add content to the document
        for index, row in df.iterrows():
            if row[p1] == 'Issue Details' or row[p1] == 'Action Plan Details':
                # Stop processing the loop, but do not exit the function
                break
            if row[p1] and not exclude_pattern.search(row[p1]):
                # Add a colon after the heading text and apply color
                heading_text = f"{row[p1]}:"
                heading = doc.add_heading(heading_text, level=2)
                for run in heading.runs:
                    run.font.color.rgb = heading_color
            if p2 and row[p2] and not exclude_pattern.search(row[p2]):
                # Add a colon after the heading text and apply color
                heading_text = f"{row[p2]}:"
                heading = doc.add_heading(heading_text, level=3)
                for run in heading.runs:
                    run.font.color.rgb = heading_color
            # Add content from p3 column
            lines = row[p3].split('\n')
            for line in lines:
                if exclude_pattern.search(line):
                    continue
                stripped_line = line.strip()
                if stripped_line.startswith('-'):
                    paragraph = doc.add_paragraph(stripped_line.lstrip('- ').strip(), style='ListBullet')
                else:
                    paragraph = doc.add_paragraph(stripped_line)
                paragraph.style.font.size = Pt(10)
                paragraph.paragraph_format.space_after = spacing_after
        
        # Add Issue Details section
        spacing_paragraph = doc.add_paragraph()
        spacing_paragraph.paragraph_format.space_before = Pt(20)  # Adjust the point value as needed

        # Add merged_dataframe to the document
        add_dataframe_to_doc(merged_dataframe, doc)
        
        # Add footer text
        footer = section.footer
        footer_paragraph = footer.add_paragraph()
        footer_text = ("This report is confidential and not to be distributed to anyone beyond the individuals "
                       "indicated. Should copies of this report be requested by any employee of Prudential or "
                       "subsidiaries, the request should be referred to the Internal Audit Department. All requests "
                       "by external parties, either individual or regulatory entity should be referred to the Law Department.")
        run = footer_paragraph.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 128)
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.italic = True
        
        # Save the Word document to a file
        doc.save(output_path)
        return True

    except KeyError as e:
        print(f"KeyError occurred: {e}")
        # Handle KeyError appropriately (e.g., log, raise custom exception, return False)
        return False
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        # Handle other exceptions as needed (e.g., log, raise custom exception, return False)
        return False

def createaudit_report(df, p1, p2, p3, header_image_path,output_path, is_map=None):
    try:
        if is_map:
            mapped_issue = mapping_df(df, p1, p2, p3)
            
            action_plan_df = mapp_action_df(df, p1, p2, p3)
            issue_df = combine_issues_by_id(mapped_issue)
            
            merged_dataframe = map_dataframes(action_plan_df, issue_df, is_map)
            
            ret_val_word = create_combined_audit_report(df, p1, p2, p3, header_image_path, merged_dataframe,output_path)
            return ret_val_word
        else:
            return 'no mapping found'
    
    except Exception as e:
        print(f"Error in creating audit report PDF: {str(e)}")
        return None


# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>> Main Function >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>

def main_function(input_path, header_image_path,output_path):
    try:
        df = pd.read_excel(input_path)
        # Apply convert_to_string function to each element in the DataFrame
        str_df = df.applymap(convert_to_string)
        p1, p2, p3 = 'Section', None, 'LLM Answer'
        # Filter rows where 'Section' contains 'Mapp'
        filtered_df = df[df['Section'].astype(str).str.contains('Mapp', case=False, na=False)]
        llm_answers = filtered_df['LLM Answer'].tolist()
        ret_structure = preprocess_llm_answer(llm_answers[0])
        is_map = ret_structure.split('\n')
        
        ret_val = createaudit_report(str_df, p1, p2, p3, header_image_path,output_path,is_map)
        return ret_val
        
    except FileNotFoundError:
        print(f"Error: File '{input_path}' not found.")
    except Exception as e:
        print(f"Error: {str(e)}")

# Example usage:

