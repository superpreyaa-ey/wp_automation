
import os
import pandas as pd
from docx import Document



def determine_highest_severity(identified_issues,severity_levels):
    # ["Low", "Medium", "High"]
    # ["Effective", "Partially Effective","Not Effective"]
    
    # Initialize the highest severity index to the lowest level
    highest_severity_index = 0 

    # Loop through each identified issue to find the highest severity
    for issue in identified_issues:
        current_severity_index = severity_levels.index(issue)
        if current_severity_index > highest_severity_index:
            highest_severity_index = current_severity_index

    return severity_levels[highest_severity_index]



def extract_components(folder_path):
    
    parts = folder_path.split('/')
    if len(parts) >= 4:
        operating_division = parts[2]
        business_group = f"{parts[3]}"
        reportable_segment = parts[4]
    else:
        operating_division = business_group = reportable_segment = "Unknown"
    return operating_division, business_group, reportable_segment



def create_issues_doc(issue_id, issue_group,output_path):
    doc = Document()
    doc.add_heading(f'Issues for ID: {issue_id}', level=1)
    
    df1_issues = issue_group.T
    df1_issues.index.name = 'old_index'
    df1_issues.reset_index(inplace=True)
    df1_issues.rename(columns={'old_index': 'heading'}, inplace=True)

    for _, row in df1_issues.iterrows():
        heading = row['heading']
        value = row.iloc[1]
        doc.add_heading(f'{heading}:', level=2)
        doc.add_paragraph(str(value))
    doc_output_path = os.path.join(output_path, f'Issue_{issue_id}.docx')
    
    doc.save(doc_output_path)
    return True
    # print(f'Document for Issue ID {issue_id} saved to {doc_output_path}')

def create_action_plans_doc(action_plans_groups,output_path):
    doc = Document()
    doc.add_heading('Consolidated Action Plans', level=1)

    counter = 1

    for action_id, action_group in action_plans_groups:
        doc.add_heading(f'Action Plan {counter}', level=2)
        
        df1_action_plans = action_group.T
        df1_action_plans.index.name = 'old_index'
        df1_action_plans.reset_index(inplace=True)
        df1_action_plans.rename(columns={'old_index': 'heading'}, inplace=True)
    
        for _, row in df1_action_plans.iterrows():
            heading = row['heading']
            value = row.iloc[1]
            doc.add_heading(f'{heading}:', level=3)
            doc.add_paragraph(str(value))

        counter += 1

    doc_output_path = os.path.join(output_path, 'Action_Plans.docx')
    doc.save(doc_output_path)
    print(f'Document saved to {doc_output_path}')
    return True
def process_groups(issues_groups, action_plans_groups,output_path):
    try:
        
        for issue_id, issue_group in issues_groups:
            create_issues_doc(issue_id, issue_group,output_path)
        create_action_plans_doc(action_plans_groups,output_path)
    except Exception as e:
        print(f"An error occurred while processing groups: {e}")
        
def process_excel_files(folder_path, columns_to_keep, columns_ap):
    all_issues = pd.DataFrame()
    all_action_plans = pd.DataFrame()

    try:        
        for filepath in folder_path:
            get_base_path = os.getcwd() + filepath
            file_path = os.path.abspath(get_base_path)

            try:
                df_issues_1 = pd.read_excel(file_path, sheet_name='Issues')
                df_issues = df_issues_1[columns_to_keep]
                df_issues.columns = df_issues.iloc[0]
                df_issues = df_issues[1:].reset_index(drop=True)
                df_issues['Operating Division'], df_issues['Business Group'], df_issues['Reportable Segment'] = zip(*df_issues['Folder Path'].apply(extract_components))

                df_action_plans_1 = pd.read_excel(file_path, sheet_name='Action Plans')
                df_action_plans = df_action_plans_1[columns_ap]
                df_action_plans.columns = df_action_plans.iloc[0]
                df_action_plans = df_action_plans[1:].reset_index(drop=True)
                ap_ids = df_action_plans['ID'].to_list()
                df_issues['Action Plans ID'] = [ap_ids] * len(df_issues)
                df_action_plans['Issue_ID'] = [df_issues['ID'].tolist()] * len(df_action_plans)

                all_issues = pd.concat([all_issues, df_issues])
                all_action_plans = pd.concat([all_action_plans, df_action_plans])

            except Exception as e:
                print(f"Error processing file {file_path}: {e}")

        issues_groups = all_issues.groupby('ID')
        action_plans_groups = all_action_plans.groupby('ID')
        identified_issues = all_issues['Issue Rating'].tolist()
        print(f" >>>>>", identified_issues)
        issue_val = determine_highest_severity(identified_issues,["Low", "Medium", "High"])
        print("Final issue", issue_val)
        return issues_groups,action_plans_groups,issue_val

    except FileNotFoundError:
        print("The folder path provided does not exist.")
    except Exception as e:
        print(f"An error occurred: {e}")
        

columns_to_keep = [
    'Folder Path',
    'System Fields.Name',
    'OPSS-Iss.Priority',
    'Pru-Iss.RptIss',
    'Pru-Shrd-RC.RtCsDesc',
    'System Fields.Description',
    'OPSS-Iss.Assignee.Display Name',
    'OPSS-Iss.Due Date',
    'Pru-Iss.Desc',
    'Pru-Iss.IssRtRat',
    'Pru-Shrd-Cat.RiskCat1',
    'OPSS-Iss.Status',
    'OPSS-Iss.Additional Description'
]

# Columns to keep for 'Action Plans'
columns_ap = [
    'System Fields.Name',
    'Pru-Action.Desc',
    'OPSS-AI.Assignee.Display Name',
    'OPSS-AI.Due Date',
    'OPSS-AI.Status'
]


def create_action_issue_doc(folder_path, columns_to_keep, columns_ap,output_path,only_issueval=None):
    try:
        issues_groups,action_plans_groups,issue_val = process_excel_files(folder_path, columns_to_keep, columns_ap)
        if only_issueval == None:
            process_groups(issues_groups, action_plans_groups,output_path)
        
        return issue_val
    except Exception as e:
        print(f"Error creating documents for action plans: {e}")

        
# Issue_var = create_action_issue_doc(folder_path, columns_to_keep, columns_ap,only_issueval=True)
# print(f">>>>> Return val {Issue_var}")

# >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>> Work Paper >>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>

# Specify the folder path


from docx.shared import Inches


def update_control_condition(Control_Dict):
    # Retrieve values from the dictionary

    design_effectiveness = Control_Dict["Design_Effectiveness"]
    operating_effectiveness = Control_Dict["Operating Effectiveness"]
    issue_rating = Control_Dict["Issue"]

    # Initialize control_val as an empty string
    control_val = ""

    """
    If both the “Design Effectiveness” and “Operating Effectiveness” are “Effective” and the “Final issue rating” is low then the “Control Condition” will be “Satisfactory with limited exceptions”.
    If both the “Design Effectiveness” and “Operating Effectiveness” are “Effective” and the “Final issue rating” is Medium then the “Control Condition” will be “Partially Satisfactory”.
    If both the “Design Effectiveness” and “Operating Effectiveness” are “Effective” and the “Final issue rating” is High then the “Control Condition” will be “Unsatisfactory”.
    If either of the “Design Effectiveness” and “Operating Effectiveness” are Not Effective and the “Final issue rating” is High, then the “Control Condition” will be “Unsatisfactory”.
    If either of the “Design Effectiveness” and “Operating Effectiveness” are Not Effective and the “Final issue rating” is Low, then the “Control Condition” will be “Partially Satisfactory”.
    If either of the “Design Effectiveness” and “Operating Effectiveness” are Not Effective and the “Final issue rating” is Medium, then the “Control Condition” will be “Unsatisfactory”.
    """
    # Check the conditions and update control_val accordingly
    if design_effectiveness == "Effective" and operating_effectiveness == "Effective":
        if issue_rating == "Low":
            control_val = "Satisfactory with limited exceptions"
        elif issue_rating == "Medium":
            control_val = "Partially Satisfactory"
        elif issue_rating == "High":
            control_val = "Unsatisfactory"

    elif design_effectiveness == "Not Effective" or operating_effectiveness == "Not Effective":
        if issue_rating == "Low":
            control_val = "Satisfactory with limited exceptions" 
        elif issue_rating == "Medium":
            control_val = "Partially Satisfactory" #"Partially Satisfactory"
        elif issue_rating == "High":
            control_val = "Unsatisfactory"
            
    elif design_effectiveness == "Not Effective" and operating_effectiveness == "Not Effective":

        if issue_rating == "Low":
            control_val = "Partially Satisfactory" #"Satisfactory with limited exceptions"
        elif issue_rating == "Medium":
            control_val = "“Unsatisfactory”" #"Partially Satisfactory"
        elif issue_rating == "High":
            control_val = "Unsatisfactory"
            
    #if design_effectiveness
    # elif design_effectiveness == "Partially Effective" or operating_effectiveness == "Not Effective": # or (design_effectiveness == "Not Effective" or operating_effectiveness == "Partially Effective"):
    #     if issue_rating == "Low":
    #         control_val = "Partially Satisfactory"#"Satisfactory with limited exceptions"
    #     elif issue_rating == "Medium":
    #         control_val = "Partially Satisfactory"
    #     elif issue_rating == "High":
    #         control_val = "Unsatisfactory"
            
    elif design_effectiveness == "Partially Effective" and operating_effectiveness == "Partially Effective":
        if issue_rating == "Low":
            control_val = "Satisfactory with limited exceptions"
        elif issue_rating == "Medium":
            control_val = "Partially Satisfactory"
        elif issue_rating == "High":
            control_val = "Unsatisfactory"
    # Update the Control Condition in the dictionary
    # Control_Dict['Control Condition'] = control_val

    # Return the updated dictionary
    return control_val



def process_issues_from_singlerow(control_list, output_folder):
    # Iterate over each dictionary in the control list
    
    for control_dict in control_list:
        # Extract values from the dictionary
        design_effectiveness = control_dict.get("Design_Effectiveness", "")
        operating_effectiveness = control_dict.get("Operationg Effectiveness", "")  # Note the typo in "Operating"
        issue = control_dict.get("Issue", "")
        
        # Create a new Word document
        doc = Document()
        
        # Create a filename based on the extracted values
        filename = f"{design_effectiveness} {operating_effectiveness}.docx"
        
        # Define the tab stop position (e.g., 3 inches)
        tab_stop_position = Inches(3)
        
        # Add paragraphs with headings and values side by side
        p = doc.add_paragraph()
        p.add_run('Design Effectiveness:\t').bold = True
        p.paragraph_format.tab_stops.add_tab_stop(tab_stop_position)
        p.add_run(str(design_effectiveness))
        
        p = doc.add_paragraph()
        p.add_run('Operating Effectiveness:\t').bold = True
        p.paragraph_format.tab_stops.add_tab_stop(tab_stop_position)
        p.add_run(str(operating_effectiveness))
        
        p = doc.add_paragraph()
        p.add_run('Issue:\t').bold = True
        p.paragraph_format.tab_stops.add_tab_stop(tab_stop_position)
        p.add_run(str(issue))
        
        # Save the document to the specified output folder
        doc.save(output_folder + "\\" + filename)


def process_issues_from_single_line(control_list, output_folder):
    # Iterate over each dictionary in the control list
    for control_dict in control_list:
        # Extract values from the dictionary
        design_effectiveness = control_dict.get("Design_Effectiveness", "")
        operating_effectiveness = control_dict.get("Operationg Effectiveness", "")  # Note the typo in "Operating"
        issue = control_dict.get("Issue", "")
        
        # Create a new Word document
        doc = Document()
        
        # Create a filename based on the extracted values
        filename = f"{design_effectiveness} {operating_effectiveness}.docx"
        
        # Add a table with headings and values side by side
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'

        # Fill in the table with headings and values
        table.cell(0, 0).text = 'Design Effectiveness:'
        table.cell(0, 1).text = str(design_effectiveness)
        table.cell(1, 0).text = 'Operating Effectiveness:'
        table.cell(1, 1).text = str(operating_effectiveness)
        table.cell(2, 0).text = 'Issue:'
        table.cell(2, 1).text = str(issue)
        
        # Save the document to the specified output folder
        doc.save(output_folder + "/" + filename)


def process_issues_from_2(control_list, output_folder,report_name='same'):
    # Iterate over each dictionary in the control list
    for control_dict in control_list:
        # Extract values from the dictionary
        design_effectiveness = control_dict.get("Design_Effectiveness", "")
        operating_effectiveness = control_dict.get("Operating Effectiveness", "")  # Note the typo in "Operating"
        issue = control_dict.get("Issue", "")
        control_condition = control_dict.get("Control Condition", "")

        
        # Create a new Word document
        doc = Document()
        
        # Create a filename based on the extracted values # output path for issue
        filename = f"{'Work Paper_'} {report_name}.docx"
        
        # Add headings and paragraphs to the document
        doc.add_heading(f'Design Effectiveness:', level=2)
        doc.add_paragraph(str(design_effectiveness))
        doc.add_heading(f'Operating Effectiveness:', level=2)
        doc.add_paragraph(str(operating_effectiveness))
        doc.add_heading(f'Final Issue:', level=2)
        doc.add_paragraph(str(issue))
        doc.add_heading(f'Control Condition:', level=2)
        doc.add_paragraph(str(control_condition))
        # Save the document to the specified output folder
        doc.save(output_folder + "\\" + filename)






def process_workpapers(folder_path, control_dict, control_lst, Issue_var):
    columns_to_keep = ['Parent Objects', 'Workpaper Result (Design Effectiveness)', 'Workpaper Result (Operating Effectiveness)']
    all_issues = pd.DataFrame()

    try:
        # for filename in os.listdir(folder_path):
        
        for filepath in folder_path:
            # if filename.endswith(".xls") or filename.endswith(".xlsx"):
            # file_path = os.path.join(folder_path, filename)
            get_base_path = os.getcwd() + filepath
            file_path = os.path.abspath(get_base_path)
            df_issues_1 = pd.read_excel(file_path, sheet_name='Workpapers')
            df_issues = df_issues_1[columns_to_keep]
            all_issues = pd.concat([all_issues, df_issues])

        WR_Design_Effectiveness = all_issues['Workpaper Result (Design Effectiveness)'].tolist()
        WR_Operating_Effectiveness = all_issues['Workpaper Result (Operating Effectiveness)'].tolist()

        WRDesign_Effectiveness = determine_highest_severity(WR_Design_Effectiveness,["Effective", "Partially Effective","Not Effective"])
        WROperating_Effectiveness = determine_highest_severity(WR_Operating_Effectiveness,["Effective", "Partially Effective","Not Effective"])

        control_dict["Design_Effectiveness"] = WRDesign_Effectiveness
        control_dict["Operating Effectiveness"] = WROperating_Effectiveness
        control_dict["Issue"] = Issue_var
        control_val = update_control_condition(control_dict)
        control_dict['Control Condition'] = control_val

        control_lst.append(control_dict)

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except pd.errors.EmptyDataError as e:
        print(f"Error: No data found in one of the files: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

    return control_lst

def get_control_condition(folder_path,Issue_var,output_path):
    try:
        Control_Dict = {}
        control_lst = []
        
        process_workpapers(folder_path, Control_Dict, control_lst, Issue_var)
        process_issues_from_2(control_lst,output_path)
        # issues_groups,action_plans_groups,issue_val = process_excel_files(folder_path, columns_to_keep, columns_ap)

        return True
    except Exception as e:
        print(f"Error creating documents for action plans: {e}")
        
# get_controllist = get_control_condition(folder_path,Issue_var)
# print(get_controllist)



# ?*******************************************************************************************************************************************************************************************************

# wp =['/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/Export-Audit_Section-13718-FMPGIM.xls', '/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/Export-Audit_Section-5335-FMPGIM.xls', '/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/Export-Audit_Section-68107-FMPGIM.xls', '/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/Export-Audit_Section-8201-FMPGIM.xls', '/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/Export-Audit_Section-90309-FMPGIM.xls']

# IS =['/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/001402 PGIM FM.xls', '/static/media/project_files/audit_check_files/shubham/IS_WP_TEST-2025/PreProcess Test/Export-Issue-653.xls']


def vaidate_wp_is_main(is_doc_path,wp_doc_path,output_path):
    try:

        issue_val = create_action_issue_doc(is_doc_path, columns_to_keep, columns_ap,output_path)
        print(f"Issue Value: {issue_val}")

        # Call the second function
        success = get_control_condition(wp_doc_path, issue_val,output_path)
        if success:
            print("Control condition processed successfully.")

    except Exception as e:
        print(f"An error occurred in the main function: {e}")

